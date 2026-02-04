"""
Word Document (DOCX) Parser Module
Uses python-docx for native DOCX handling
"""

import re
import os
from pathlib import Path
from typing import Dict, List, Any, Optional
from dataclasses import dataclass, field
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.table import Table as DocxTable


@dataclass
class TextBlock:
    """Represents a block of text with formatting info"""
    text: str
    font_size: float
    font_name: str
    is_bold: bool
    is_italic: bool
    bbox: tuple  # (x0, y0, x1, y1) - not applicable for DOCX, kept for compatibility
    page_num: int
    style_name: Optional[str] = None
    is_heading: bool = False
    heading_level: int = 0


@dataclass
class ImageInfo:
    """Represents an extracted image"""
    image_bytes: bytes
    width: int
    height: int
    page_num: int
    bbox: tuple
    caption: Optional[str] = None
    filename: Optional[str] = None


@dataclass
class TableInfo:
    """Represents an extracted table"""
    rows: List[List[str]]
    page_num: int
    bbox: tuple
    caption: Optional[str] = None


@dataclass
class ParsedDocument:
    """Complete parsed document structure"""
    text_blocks: List[TextBlock] = field(default_factory=list)
    images: List[ImageInfo] = field(default_factory=list)
    tables: List[TableInfo] = field(default_factory=list)
    full_text: str = ""
    page_count: int = 0
    metadata: Dict[str, Any] = field(default_factory=dict)


class DOCXParser:
    """
    Word document parser using python-docx
    Extracts text with styles, images, and tables
    """
    
    def __init__(self, upload_dir: str = "uploads"):
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(exist_ok=True)
    
    def parse(self, docx_path: str) -> ParsedDocument:
        """
        Parse a DOCX file and extract all content
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            ParsedDocument with extracted content
        """
        doc = Document(docx_path)
        result = ParsedDocument()
        
        # Extract metadata
        result.metadata = self._extract_metadata(doc)
        
        full_text_parts = []
        current_page = 0  # DOCX doesn't have explicit pages
        
        # Process paragraphs
        for para in doc.paragraphs:
            text_block = self._process_paragraph(para, current_page)
            if text_block:
                result.text_blocks.append(text_block)
                full_text_parts.append(para.text)
        
        # Process tables
        for table_idx, table in enumerate(doc.tables):
            table_info = self._process_table(table, current_page, table_idx)
            if table_info:
                result.tables.append(table_info)
        
        # Extract images
        images = self._extract_images(doc, docx_path)
        result.images = images
        
        result.full_text = "\n".join(full_text_parts)
        result.page_count = max(1, len(full_text_parts) // 40)  # Rough estimate
        
        return result
    
    def _extract_metadata(self, doc: Document) -> Dict[str, Any]:
        """Extract document metadata"""
        core_props = doc.core_properties
        return {
            "author": core_props.author or "",
            "title": core_props.title or "",
            "subject": core_props.subject or "",
            "keywords": core_props.keywords or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
        }
    
    def _process_paragraph(self, para, page_num: int) -> Optional[TextBlock]:
        """Process a single paragraph and extract formatting"""
        text = para.text.strip()
        if not text:
            return None
        
        # Determine if it's a heading
        style_name = para.style.name if para.style else ""
        is_heading = "Heading" in style_name or "Title" in style_name
        heading_level = 0
        
        if is_heading:
            # Extract heading level from style name (e.g., "Heading 1" -> 1)
            match = re.search(r'\d+', style_name)
            if match:
                heading_level = int(match.group())
        
        # Get formatting from first run (simplified - assumes consistent formatting)
        font_size = 12.0
        font_name = "Times New Roman"
        is_bold = False
        is_italic = False
        
        if para.runs:
            run = para.runs[0]
            if run.font.size:
                font_size = run.font.size.pt
            if run.font.name:
                font_name = run.font.name
            is_bold = run.bold or False
            is_italic = run.italic or False
        
        return TextBlock(
            text=text,
            font_size=font_size,
            font_name=font_name,
            is_bold=is_bold,
            is_italic=is_italic,
            bbox=(0, 0, 0, 0),
            page_num=page_num,
            style_name=style_name,
            is_heading=is_heading,
            heading_level=heading_level
        )
    
    def _process_table(self, table: DocxTable, page_num: int, table_idx: int) -> Optional[TableInfo]:
        """Extract table content and structure"""
        rows = []
        
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_cells.append(cell_text)
            rows.append(row_cells)
        
        if not rows:
            return None
        
        # Try to find caption (look for "Table X" pattern in nearby paragraphs)
        caption = f"Table {table_idx + 1}"
        
        return TableInfo(
            rows=rows,
            page_num=page_num,
            bbox=(0, 0, 0, 0),
            caption=caption
        )
    
    def _extract_images(self, doc: Document, docx_path: str) -> List[ImageInfo]:
        """Extract embedded images from the DOCX file"""
        images = []
        job_id = Path(docx_path).stem
        image_dir = self.upload_dir / f"{job_id}_images"
        image_dir.mkdir(exist_ok=True)
        
        # Access document's image relationships
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    image_data = rel.target_part.blob
                    
                    # Determine image extension
                    content_type = rel.target_part.content_type
                    ext = "png"
                    if "jpeg" in content_type or "jpg" in content_type:
                        ext = "jpg"
                    elif "png" in content_type:
                        ext = "png"
                    elif "gif" in content_type:
                        ext = "gif"
                    
                    # Save image
                    img_index = len(images) + 1
                    filename = f"image_{img_index}.{ext}"
                    image_path = image_dir / filename
                    
                    with open(image_path, "wb") as f:
                        f.write(image_data)
                    
                    # Get image dimensions (would need PIL for accurate dimensions)
                    images.append(ImageInfo(
                        image_bytes=image_data,
                        width=0,  # Will be determined during processing
                        height=0,
                        page_num=0,
                        bbox=(0, 0, 0, 0),
                        caption=f"Figure {img_index}",
                        filename=str(image_path)
                    ))
                except Exception as e:
                    print(f"Error extracting image: {e}")
                    continue
        
        return images
